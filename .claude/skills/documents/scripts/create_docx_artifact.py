#!/usr/bin/env python3
"""Create a small, valid DOCX artifact from stdin.

This deterministic helper covers the common simple "write this as a Word
document" path for the documents skill. It intentionally writes a real .docx
file in the workspace so the Artifact tool can publish it via file_path.
"""

from __future__ import annotations

import os
import re
import sys
import time
from pathlib import Path

from docx import Document
from docx.shared import Pt


DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


def main() -> int:
    raw = sys.stdin.read()
    text = normalize_input(raw)
    if not text:
        print("skill_error: No document body was provided.")
        return 1
    if looks_like_generation_brief(text):
        print(
            "skill_error: requires_final_body: input looks like a generation brief. "
            "Compose the final document body first, then pass only that body to this helper."
        )
        return 1

    workspace = Path(os.environ.get("AGENT_WORKSPACE_DIR") or os.getcwd())
    output_dir = workspace / "generated-artifacts"
    output_dir.mkdir(parents=True, exist_ok=True)

    title, body = extract_title(text)
    requested = os.environ.get("DOCX_FILENAME", "").strip()
    filename = unique_filename(output_dir, safe_filename(requested or title) + ".docx")
    output_path = output_dir / filename
    write_docx(output_path, title, split_paragraphs(body))

    print(f"output_file: {output_path}")
    print(f"artifact_file_path: generated-artifacts/{filename}")
    print(f"filename: {filename}")
    print(f"content_type: {DOCX_CONTENT_TYPE}")
    return 0


def normalize_input(raw: str) -> str:
    text = raw.replace("\r\n", "\n").replace("\r", "\n").strip()
    text = re.sub(r"^\s*/documents\b", "", text, flags=re.IGNORECASE).strip()
    return text


def looks_like_generation_brief(text: str) -> bool:
    compact = re.sub(r"\s+", "", text).lower()
    if len(compact) < 40:
        return False
    request_markers = (
        "请生成",
        "帮我生成",
        "生成一个",
        "生成一份",
        "createa",
        "writea",
    )
    body_markers = (
        "内容包括",
        "文件名为",
        "要求",
        "format",
        "filename",
    )
    return any(marker in compact for marker in request_markers) and any(
        marker in compact for marker in body_markers
    )


def extract_title(text: str) -> tuple[str, str]:
    lines = [line.rstrip() for line in text.splitlines()]
    for index, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        heading = re.match(r"^#{1,2}\s+(.+)$", stripped)
        if heading:
            title = clean_title(heading.group(1))
            body = "\n".join(lines[:index] + lines[index + 1 :]).strip()
            return title, body or title
        if len(stripped) <= 80 and not sentence_like(stripped):
            title = clean_title(stripped)
            body = "\n".join(lines[:index] + lines[index + 1 :]).strip()
            return title, body or title
        break
    return "Word Document", text


def split_paragraphs(text: str) -> list[str]:
    text = text.strip()
    if "\n" not in text:
        text = re.sub(r"([.!?。！？])\s+", r"\1\n", text)
    paragraphs = [line.strip() for line in re.split(r"\n+", text) if line.strip()]
    return paragraphs or [text]


def write_docx(path: Path, title: str, paragraphs: list[str]) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    heading = doc.add_heading(title, level=0)
    for run in heading.runs:
        run.font.name = "Arial"

    for item in paragraphs:
        para = doc.add_paragraph(item)
        for run in para.runs:
            run.font.name = "Arial"

    doc.save(path)


def safe_filename(name: str) -> str:
    base = Path(name).stem if name else ""
    base = base.strip() or "generated-document"
    base = re.sub(r"[\\/:*?\"<>|]+", "_", base)
    base = re.sub(r"\s+", "_", base)
    base = base.strip("._") or "generated-document"
    return base[:80]


def unique_filename(output_dir: Path, filename: str) -> str:
    stem = Path(filename).stem
    suffix = Path(filename).suffix or ".docx"
    candidate = stem + suffix
    if not (output_dir / candidate).exists():
        return candidate
    stamp = time.strftime("%Y%m%dT%H%M%SZ", time.gmtime())
    return f"{stem}-{stamp}{suffix}"


def clean_title(text: str) -> str:
    title = re.sub(r"\s+", " ", text.strip(" #\t:："))
    return title[:80] or "Word Document"


def sentence_like(text: str) -> bool:
    return bool(re.search(r"[。.!?？；;，,]", text))


if __name__ == "__main__":
    raise SystemExit(main())
